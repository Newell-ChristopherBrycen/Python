#!python3

import docx

'''d = docx.Document('c:\\users\\bryne\\documents\\TestDoc.docx')

print(d.paragraphs[0].text)

#print(d.paragraphs[2].text)

#print(d.paragraphs[1].text)


p = d.paragraphs[0]

print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
print(p.runs[4].text)
print(p.runs[5].text)

print(str(p.runs[4].bold == None))
p.runs[5].text = 'Python3.'
       
p.runs[5].italic = True
       
p.runs[5].bold = True
       
p.runs[5].underline = True

p.style = 'Title'
       
d.save('c:\\users\\bryne\\documents\\TestDoc.docx')
'''

d = docx.Document() #creates a new doc

d.add_paragraph('Hello World!')

d.add_paragraph('This is paragraph 2')

d.save('c:\\users\\bryne\\documents\\NewDoc.docx')

p = d.paragraphs[0]
p.add_run('This is a new Run')
p.runs[1].bold = True
d.save('c:\\users\\bryne\\documents\\NewDoc.docx')

